"""
Investment Advice Note — PDF & DOCX Report Generator
─────────────────────────────────────────────────────
Generates SEBI-compliant Investment Advice Notes in PDF and DOCX formats.
Matches the layout defined in docs/SEBI_Investment_Advice_Note_Sample.md.

Uses:
  - fpdf2 (BaseReportPDF) for PDF generation
  - python-docx for Word document generation
"""
import io
import os
import json
from datetime import datetime, timedelta
from typing import Optional, List, Dict, Any

from app.utils.pdf_generator import BaseReportPDF

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# ── Color Palette ──────────────────────────────────────────────────
_NAVY = (0, 31, 63)
_ACCENT_BLUE = (0, 70, 160)
_LIGHT_BLUE_BG = (235, 244, 255)
_SECTION_BG = (245, 246, 248)
_TABLE_HEADER_BG = (220, 232, 246)
_BORDER = (200, 210, 220)
_TEXT_DARK = (20, 20, 20)
_TEXT_MUTED = (110, 110, 110)
_ROW_ALT = (250, 251, 253)
_GREEN = (0, 120, 50)
_GREEN_BG = (235, 248, 235)
_RED_MUTED = (180, 60, 60)

# ── SEBI Disclaimer ───────────────────────────────────────────────
_SEBI_DISCLAIMER = (
    "Investment in securities is subject to market risks. This Investment Advice Note is prepared "
    "on the basis of information which the adviser considers reliable but does not represent or "
    "warrant its accuracy or completeness. This document does not constitute an offer or solicitation "
    "to buy or sell any security. Past performance is not indicative of future returns. The Investment "
    "Adviser does not guarantee any return and shall not incur any liability by reason of any loss "
    "which a client may suffer due to any depletion in the value of assets under advice, fluctuation "
    "in asset value, non-performance or underperformance of securities or funds or other market "
    "conditions. This note is issued exclusively for the addressee client and shall not be reproduced "
    "or redistributed. Registration of Investment Adviser does not guarantee quality of advice."
)

_RECORD_RETENTION = (
    "Record retention: 5 years from date of issue  |  Both IA and client must retain this document  |  "
    "Regulation 22, SEBI (Investment Advisers) Regulations 2013 as amended"
)


def format_indian_number(val):
    if val is None:
        return ""
    try:
        val_float = float(val)
        if val_float.is_integer():
            int_part = str(int(val_float))
            dec_part = ""
        else:
            s = f"{val_float:.2f}"
            int_part, dec_part = s.split('.')
            dec_part = "." + dec_part
            
        if len(int_part) <= 3:
            grouped = int_part
        else:
            last_three = int_part[-3:]
            rest = int_part[:-3]
            rest_groups = []
            while len(rest) > 2:
                rest_groups.insert(0, rest[-2:])
                rest = rest[:-2]
            if rest:
                rest_groups.insert(0, rest)
            grouped = ",".join(rest_groups) + "," + last_three
            
        return grouped + dec_part
    except Exception:
        return str(val)


def format_amount_units_python(rec):
    ttype = rec.get("transaction_type")
    if not ttype:
        return rec.get("amount_units") or ""
        
    freq = rec.get("frequency")
    amount = rec.get("amount")
    custom_inst = rec.get("custom_instruction")
    p_type = rec.get("product_type", "")
    
    is_life_insurance = False
    if p_type:
        is_life_insurance = p_type.lower() in ("life-insurance", "life_insurance")
        
    if ttype == 'HOLDING':
        return "Existing holding"
    elif ttype == 'TEXT_ONLY':
        return custom_inst or ""
    
    formatted_amount = format_indian_number(amount) if amount is not None else ""
    
    if ttype == 'LUMP_SUM':
        return f"Rs. {formatted_amount} lump sum"
    if ttype == 'SWITCH_IN':
        return f"Rs. {formatted_amount} Switch In"
    if ttype == 'SWITCH_OUT':
        return f"Rs. {formatted_amount} Switch Out"
    if ttype == 'TRANSFER_IN':
        return f"Rs. {formatted_amount} Transfer In"
    if ttype == 'TRANSFER_OUT':
        return f"Rs. {formatted_amount} Transfer Out"
        
    freq_label = ""
    if freq == 'MONTHLY':
        freq_label = "month"
    elif freq == 'QUARTERLY':
        freq_label = "quarter"
    elif freq == 'HALF_YEARLY':
        freq_label = "half-year"
    elif freq == 'YEARLY':
        freq_label = "year"
        
    if is_life_insurance and freq == 'YEARLY':
        return f"Annual prem. Rs. {formatted_amount}"
        
    if ttype in ('SIP', 'STP', 'SWP'):
        return f"Rs. {formatted_amount}/{freq_label} {ttype}"
        
    return rec.get("amount_units") or ""


def calculate_age(dob_str):
    if not dob_str:
        return None
    try:
        import re
        parts = re.findall(r'\d+', dob_str)
        if len(parts) >= 3:
            # Check if first part is year (4 digits)
            if len(parts[0]) == 4:
                year = int(parts[0])
            elif len(parts[2]) == 4:
                year = int(parts[2])
            else:
                return None
            return 2026 - year
    except Exception:
        pass
    return None


def format_dob(dob_str):
    if not dob_str:
        return "N/A"
    try:
        dt = datetime.strptime(dob_str.split('T')[0], "%Y-%m-%d")
        return dt.strftime("%d %B %Y")
    except Exception:
        return dob_str


def format_date_issue(date_str):
    if not date_str:
        return "N/A"
    try:
        dt = datetime.strptime(date_str.split('T')[0], "%Y-%m-%d")
        return dt.strftime("%d %B %Y")
    except Exception:
        return date_str


class InvestmentAdviceNotePDF:
    """Generates a SEBI-compliant Investment Advice Note PDF using fpdf2."""

    @staticmethod
    def _render_cover_page(
        pdf: BaseReportPDF,
        note_data: dict,
        ia_data: Optional[dict],
        client: dict,
        logo_path: Optional[str] = None,
    ):
        pdf.add_page()

        # Border
        pdf.set_draw_color(*_ACCENT_BLUE)
        pdf.set_line_width(0.5)
        pdf.rect(5, 5, 200, 287)
        pdf.set_line_width(0.2)

        # Logo
        pdf.set_y(35)
        if logo_path and os.path.exists(logo_path):
            pdf.image(logo_path, 85, pdf.get_y(), 40)
            pdf.set_y(pdf.get_y() + 50)
        else:
            pdf.set_y(60)

        # Entity name
        entity_name = ""
        if ia_data:
            entity_name = ia_data.get("name_of_entity") or ia_data.get("name_of_ia", "")
        if entity_name:
            pdf.set_font("helvetica", "B", 13)
            pdf.set_text_color(*_TEXT_DARK)
            pdf.cell(0, 8, entity_name.upper(), ln=True, align="C")

        # Report title
        pdf.ln(20)
        pdf.set_font("helvetica", "B", 24)
        pdf.set_text_color(*_ACCENT_BLUE)
        pdf.cell(0, 14, "INVESTMENT ADVICE NOTE", ln=True, align="C")

        # Subtitle
        pdf.set_font("helvetica", "", 9)
        pdf.set_text_color(*_TEXT_MUTED)
        pdf.cell(0, 6, "Issued under SEBI (Investment Advisers) Regulations, 2013", ln=True, align="C")
        pdf.cell(0, 5, "Regulation 16 & 17  |  Master Circular dated 21 May 2024", ln=True, align="C")

        # Decorative bar
        pdf.ln(4)
        pdf.set_fill_color(*_ACCENT_BLUE)
        pdf.set_xy(75, pdf.get_y())
        pdf.cell(60, 1.5, "", fill=True, ln=True)

        # Client info
        pdf.ln(25)
        pdf.set_font("helvetica", "B", 16)
        pdf.set_text_color(*_TEXT_DARK)
        client_name = client.get("client_name", "Client").upper()
        pdf.cell(0, 10, client_name, ln=True, align="C")

        pdf.set_font("helvetica", "B", 11)
        pdf.set_text_color(*_TEXT_MUTED)
        client_code = client.get("client_code", "N/A")
        pdf.cell(0, 6, f"CLIENT CODE: {client_code}", ln=True, align="C")

        # Advice Note Number badge
        pdf.ln(8)
        pdf.set_font("helvetica", "B", 10)
        pdf.set_text_color(*_GREEN)
        advice_no = note_data.get("advice_note_no", "N/A")
        pdf.cell(0, 6, f"ADVICE NOTE: {advice_no}", ln=True, align="C")

        # Footer
        pdf.set_y(245)
        pdf.set_font("helvetica", "I", 9)
        pdf.set_text_color(*_TEXT_MUTED)
        date_str = note_data.get("date_of_issue", datetime.now().strftime("%d %B %Y"))
        pdf.cell(0, 6, f"Date of Issue: {date_str}", ln=True, align="C")

        if ia_data:
            reg_no = ia_data.get("ia_registration_number", "")
            if reg_no:
                pdf.set_font("helvetica", "", 9)
                pdf.cell(0, 6, f"SEBI Registration No: {reg_no}", ln=True, align="C")

        pdf.set_font("helvetica", "I", 8)
        pdf.set_text_color(*_RED_MUTED)
        pdf.cell(0, 6, "CONFIDENTIAL - FOR CLIENT AND IA RECORDS ONLY", ln=True, align="C")

    @staticmethod
    def _section_header(pdf: BaseReportPDF, title: str):
        """Render a styled section header bar."""
        if pdf.get_y() > 255:
            pdf.add_page()

        pdf.set_fill_color(*_SECTION_BG)
        pdf.rect(10, pdf.get_y(), 190, 9, "F")
        pdf.set_draw_color(*_BORDER)
        pdf.set_line_width(0.3)
        pdf.rect(10, pdf.get_y(), 190, 9, "D")
        pdf.set_fill_color(*_ACCENT_BLUE)
        pdf.rect(10, pdf.get_y(), 3, 9, "F")

        pdf.set_xy(16, pdf.get_y() + 2)
        pdf.set_font("helvetica", "B", 10)
        pdf.set_text_color(*_NAVY)
        pdf.cell(0, 5, title.upper(), ln=True)
        pdf.ln(4)

    @staticmethod
    def _kv_row(pdf: BaseReportPDF, label: str, value: str, label_w: int = 42, val_w: int = 53):
        """Render a label:value pair in a 2-column grid row with dynamic font size scaling to prevent overflow."""
        pdf.set_font("helvetica", "B", 8)
        pdf.set_text_color(*_TEXT_MUTED)
        pdf.set_draw_color(*_BORDER)
        pdf.cell(label_w, 7, f" {label}", border="LBT")
        
        pdf.set_font("helvetica", "", 9)
        pdf.set_text_color(*_TEXT_DARK)
        
        val_str = f" {value or 'N/A'}"
        
        # Calculate text width and scale down font size if it overflows the cell width (minus padding)
        font_size = 9.0
        width = pdf.get_string_width(val_str)
        max_allowed_w = val_w - 3  # 3mm safety margin for padding
        
        while width > max_allowed_w and font_size > 6.0:
            font_size -= 0.5
            pdf.set_font("helvetica", "", font_size)
            width = pdf.get_string_width(val_str)
            
        # If it still overflows at 6.0pt, truncate with ellipsis
        if width > max_allowed_w:
            while len(val_str) > 4 and width > max_allowed_w:
                val_str = val_str[:-4] + "..."
                width = pdf.get_string_width(val_str)
                
        pdf.cell(val_w, 7, val_str, border="RBT")

    @staticmethod
    def _kv_grid(pdf: BaseReportPDF, fields: List[tuple]):
        """Render key-value pairs in a 2x2 grid layout."""
        for i in range(0, len(fields), 2):
            pdf.set_x(10)
            f1 = fields[i]
            InvestmentAdviceNotePDF._kv_row(pdf, f1[0], f1[1])
            if i + 1 < len(fields):
                f2 = fields[i + 1]
                InvestmentAdviceNotePDF._kv_row(pdf, f2[0], f2[1])
            pdf.ln()

    @staticmethod
    def _full_width_text(pdf: BaseReportPDF, label: str, value: str):
        """Render a full-width label + multi-line text block."""
        pdf.set_x(10)
        pdf.set_font("helvetica", "B", 8)
        pdf.set_text_color(*_TEXT_MUTED)
        pdf.set_draw_color(*_BORDER)
        pdf.cell(42, 7, f" {label}", border="LTB")
        pdf.set_font("helvetica", "", 8)
        pdf.set_text_color(*_TEXT_DARK)
        pdf.multi_cell(148, 7, f" {value or 'N/A'}", border="RTB")

    @staticmethod
    def generate_pdf(
        note_data: dict,
        ia_data: Optional[dict] = None,
        logo_path: Optional[str] = None,
    ) -> bytes:
        """
        Generate a complete SEBI Investment Advice Note PDF.
        
        Args:
            note_data: Full advice note dict from Bridge (includes client_snapshot, recommendations)
            ia_data: IA Master data dict
            logo_path: Absolute path to the IA logo file
        """
        advisor_name = ia_data.get("name_of_ia", "") if ia_data else ""
        entity_name = ia_data.get("name_of_entity", "") if ia_data else ""
        ia_reg_no = ia_data.get("ia_registration_number", "") if ia_data else ""

        pdf = BaseReportPDF(
            advisor_name=advisor_name,
            entity_name=entity_name,
            ia_reg_no=ia_reg_no,
            header_text="Investment Advice Note - Confidential",
        )

        client = note_data.get("client_snapshot", {})
        recommendations = note_data.get("recommendations", [])

        now_ist = datetime.utcnow() + timedelta(hours=5, minutes=30)
        generated_on = now_ist.strftime("%d %b %Y, %I:%M %p")

        # ══════════════════ COVER PAGE ══════════════════
        InvestmentAdviceNotePDF._render_cover_page(pdf, note_data, ia_data, client, logo_path)

        # ══════════════════ SECTION A - IA Details ══════════════════
        pdf.add_page()
        InvestmentAdviceNotePDF._section_header(pdf, "Section A - Investment Adviser Details")

        basl_id = ia_data.get("basl_membership_id", "") if ia_data else ""
        website_val = ia_data.get("website", "") if ia_data else ""

        ia_fields = [
            ("IA / Firm Name", entity_name or advisor_name),
            ("Advice Note No.", note_data.get("advice_note_no", "N/A")),
            ("SEBI Reg. No.", ia_reg_no),
            ("Date of Issue", note_data.get("date_of_issue", "N/A")),
            ("IAASB", f"BSE Limited (IAASB) - {basl_id}" if basl_id else "BSE Limited (IAASB)"),
            ("Advice Validity", note_data.get("advice_validity_custom_text", "N/A")),
            ("Principal Officer", note_data.get("principal_officer_name", "N/A")),
            ("PO Reg. No.", note_data.get("principal_officer_reg_no", "N/A")),
            ("Website", website_val or "N/A"),
            ("Advice Category", note_data.get("advice_category", "Comprehensive Advisory")),
        ]
        InvestmentAdviceNotePDF._kv_grid(pdf, ia_fields)

        # Registered address (full width)
        if ia_data:
            addr = ia_data.get("registered_address", "")
            if addr:
                InvestmentAdviceNotePDF._full_width_text(pdf, "Registered Address", addr)
        pdf.ln(6)

        # ══════════════════ SECTION B - Client Details ══════════════════
        InvestmentAdviceNotePDF._section_header(pdf, "Section B - Client Details and Risk Profile")

        risk_score = client.get("risk_profile_score")
        risk_profile_str = client.get("risk_profile", "N/A")
        if risk_score and risk_score != "N/A":
            risk_profile_str = f"{risk_profile_str} (Score: {risk_score}/100)"

        liabilities_val = client.get("existing_liabilities")
        if liabilities_val is not None:
            try:
                liabilities_str = f"Rs. {float(liabilities_val):,.0f}"
            except Exception:
                liabilities_str = str(liabilities_val)
        else:
            liabilities_str = "N/A"

        client_fields = [
            ("Client Full Name", client.get("client_name", "N/A")),
            ("Client ID", client.get("client_code", "N/A")),
            ("PAN Number", client.get("pan_number", "N/A")),
            ("Client DOB", client.get("date_of_birth", "N/A")),
            ("Email", client.get("email", "N/A")),
            ("Mobile", client.get("phone_number", "N/A")),
            ("Risk Profile", risk_profile_str),
            ("Risk Profile Date", client.get("risk_profile_date", "N/A")),
            ("Investment Horizon", client.get("investment_horizon", "N/A")),
            ("Annual Income Band", note_data.get("annual_income_band", "N/A")),
            ("Existing Liabilities", liabilities_str),
            ("Assets Under Advice", f"Rs. {float(note_data.get('assets_under_advice', 0)):,.0f}"),
            ("Fee Mode", note_data.get("fee_mode", "N/A").replace("_", " ").title()),
            ("Fee Amount", f"Rs. {float(note_data.get('fee_amount', 0)):,.0f}"),
        ]
        InvestmentAdviceNotePDF._kv_grid(pdf, client_fields)
        
        # Address & Primary Financial Goal as full-width multi-line text boxes to prevent grid collision
        InvestmentAdviceNotePDF._full_width_text(pdf, "Address", client.get("address", "N/A"))
        InvestmentAdviceNotePDF._full_width_text(
            pdf, "Primary Financial Goal", note_data.get("primary_financial_goal", "N/A")
        )

        # Recommended Asset Allocation
        rec_alloc = note_data.get("recommended_asset_allocation")
        alloc_text = "N/A"
        sub_allocs = {}
        if rec_alloc:
            if isinstance(rec_alloc, str):
                try:
                    import json as _json
                    rec_alloc = _json.loads(rec_alloc)
                except Exception:
                    pass
            if isinstance(rec_alloc, dict):
                sub_allocs = rec_alloc.get("sub_assets") or {}
                alloc_parts = []
                for k in ("Debt", "Equity", "Commodities"):
                    if k in rec_alloc:
                        alloc_parts.append(f"{k}: {rec_alloc[k]}%")
                for k, v in rec_alloc.items():
                    if k not in ("Debt", "Equity", "Commodities", "sub_assets"):
                        alloc_parts.append(f"{k}: {v}%")
                alloc_text = "  |  ".join(alloc_parts) if alloc_parts else "N/A"
            elif isinstance(rec_alloc, str) and rec_alloc != "{}":
                alloc_text = rec_alloc

        InvestmentAdviceNotePDF._full_width_text(pdf, "Recommended Asset Allocation", alloc_text)

        # Render sub-asset allocations if present
        if sub_allocs:
            # Group sub-assets by category
            equity_map = {
                "stocks_percentage": "Direct Equity (Stocks)",
                "mutual_fund_equity_percentage": "Equity Mutual Funds",
                "ulip_equity_percentage": "Equity ULIPs",
                "etf_equity_percentage": "Equity ETFs"
            }
            debt_map = {
                "fixed_deposits_bonds_percentage": "Fixed Deposits / Bonds",
                "mutual_fund_debt_percentage": "Debt Mutual Funds",
                "ulip_debt_percentage": "Debt ULIPs",
                "etf_debt_percentage": "Debt ETFs"
            }
            comm_map = {
                "gold_etf_percentage": "Gold ETFs",
                "silver_etf_percentage": "Silver ETFs",
                "etf_commodity_percentage": "Commodity ETFs"
            }

            eq_list = []
            dt_list = []
            cm_list = []

            for k, val in sub_allocs.items():
                try:
                    num_val = float(val) if val is not None else 0.0
                except Exception:
                    num_val = 0.0
                if num_val > 0:
                    pct_str = f"{num_val:.1f}%"
                    if k in equity_map:
                        eq_list.append((equity_map[k], pct_str))
                    elif k in debt_map:
                        dt_list.append((debt_map[k], pct_str))
                    elif k in comm_map:
                        cm_list.append((comm_map[k], pct_str))

            if eq_list or dt_list or cm_list:
                pdf.set_x(10)
                pdf.set_font("helvetica", "B", 8)
                pdf.set_text_color(*_TEXT_MUTED)
                pdf.set_draw_color(*_BORDER)
                pdf.cell(190, 7, " Recommended Sub-Asset Allocation Breakdown", border="LRT", ln=True)

                # Column Headers
                pdf.set_x(10)
                pdf.set_fill_color(*_TABLE_HEADER_BG)
                pdf.set_font("helvetica", "B", 8)
                pdf.set_text_color(*_NAVY)
                pdf.cell(63, 6, "  Equity Sub-Assets", border="1", fill=True, align="L")
                pdf.cell(63, 6, "  Debt Sub-Assets", border="1", fill=True, align="L")
                pdf.cell(64, 6, "  Commodities Sub-Assets", border="1", fill=True, align="L")
                pdf.ln()

                # Print Rows
                max_len = max(len(eq_list), len(dt_list), len(cm_list))
                pdf.set_font("helvetica", "", 7.5) # Compact font size
                pdf.set_text_color(*_TEXT_DARK)

                for idx in range(max_len):
                    pdf.set_x(10)
                    fill = idx % 2 == 1
                    if fill:
                        pdf.set_fill_color(*_ROW_ALT)

                    # 1. Equity Column (width: 47 name, 16 percentage)
                    if idx < len(eq_list):
                        name, pct = eq_list[idx]
                        pdf.cell(47, 6, f"  {name}", border="1", fill=fill)
                        pdf.cell(16, 6, f"{pct} ", border="1", fill=fill, align="R")
                    else:
                        pdf.cell(47, 6, "", border="1", fill=fill)
                        pdf.cell(16, 6, "", border="1", fill=fill)

                    # 2. Debt Column (width: 47 name, 16 percentage)
                    if idx < len(dt_list):
                        name, pct = dt_list[idx]
                        pdf.cell(47, 6, f"  {name}", border="1", fill=fill)
                        pdf.cell(16, 6, f"{pct} ", border="1", fill=fill, align="R")
                    else:
                        pdf.cell(47, 6, "", border="1", fill=fill)
                        pdf.cell(16, 6, "", border="1", fill=fill)

                    # 3. Commodities Column (width: 48 name, 16 percentage)
                    if idx < len(cm_list):
                        name, pct = cm_list[idx]
                        pdf.cell(48, 6, f"  {name}", border="1", fill=fill)
                        pdf.cell(16, 6, f"{pct} ", border="1", fill=fill, align="R")
                    else:
                        pdf.cell(48, 6, "", border="1", fill=fill)
                        pdf.cell(16, 6, "", border="1", fill=fill)

                    pdf.ln()
                pdf.ln(1)

        date_of_alloc = note_data.get("date_of_allocation", "N/A")
        if date_of_alloc and date_of_alloc != "N/A":
            pdf.set_x(10)
            InvestmentAdviceNotePDF._kv_row(pdf, "Date of Allocation", str(date_of_alloc), 42, 148)
            pdf.ln()
        pdf.ln(6)

        # ══════════════════ SECTION C - Suitability ══════════════════
        InvestmentAdviceNotePDF._section_header(pdf, "Section C - Suitability Assessment [Regulation 17]")

        InvestmentAdviceNotePDF._full_width_text(
            pdf, "Advice Suitable?", note_data.get("suitability_assessment", "N/A")
        )
        InvestmentAdviceNotePDF._full_width_text(
            pdf, "Suitability Basis", note_data.get("suitability_basis", "N/A")
        )
        InvestmentAdviceNotePDF._full_width_text(
            pdf, "Investor Advice", note_data.get("investor_advice", "N/A")
        )

        # Current Asset Allocation & Rebalancing (manual text areas)
        InvestmentAdviceNotePDF._full_width_text(
            pdf, "Current Allocation", note_data.get("current_asset_allocation", "N/A")
        )
        InvestmentAdviceNotePDF._full_width_text(
            pdf, "Rebalancing Rationale", note_data.get("rebalancing_rationale", "N/A")
        )
        pdf.ln(6)

        # ══════════════════ SECTION D - Recommendations ══════════════════
        InvestmentAdviceNotePDF._section_header(pdf, "Section D - Investment Recommendations")

        if recommendations:
            # Table header
            cols = [8, 52, 30, 30, 16, 28, 26]
            headers = ["#", "Product / Scheme Name", "ISIN / Scheme Code", "Product Type", "Action", "Amount / Units", "Price / NAV"]

            pdf.set_fill_color(*_NAVY)
            pdf.set_font("helvetica", "B", 7)
            pdf.set_text_color(255, 255, 255)
            pdf.set_x(10)
            for h, w in zip(headers, cols):
                pdf.cell(w, 8, f" {h}", border=1, fill=True)
            pdf.ln()

            # Data rows
            h_unit = 4.5
            for idx, rec in enumerate(recommendations):
                cell_texts = [
                    str(idx + 1),
                    rec.get("product_name", "N/A"),
                    rec.get("isin_code_scheme_code_uin", ""),
                    rec.get("product_type", "").replace("_", " ").title(),
                    rec.get("action", "BUY"),
                    format_amount_units_python(rec),
                    f"Rs. {float(rec['indicative_price_nav']):,.2f}" if rec.get("indicative_price_nav") else "N/A",
                ]

                lines_per_col = [
                    max(len(pdf.multi_cell(w, h_unit, str(t), split_only=True)), 1)
                    for t, w in zip(cell_texts, cols)
                ]
                row_h = max(max(lines_per_col) * h_unit, 8)

                if pdf.get_y() + row_h > 272:
                    pdf.add_page()

                fill_color = _ROW_ALT if idx % 2 == 1 else (255, 255, 255)
                pdf.set_font("helvetica", "", 7)
                pdf.set_text_color(*_TEXT_DARK)

                row_x, row_y = 10, pdf.get_y()
                for t, w in zip(cell_texts, cols):
                    pdf.set_fill_color(*fill_color)
                    pdf.rect(row_x, row_y, w, row_h, "F")
                    pdf.rect(row_x, row_y, w, row_h, "D")
                    pdf.set_xy(row_x + 1, row_y + 1)
                    pdf.multi_cell(w - 2, h_unit, str(t), align="L")
                    row_x += w
                pdf.set_y(row_y + row_h)

            # Indicative price note
            pdf.ln(2)
            pdf.set_font("helvetica", "I", 6.5)
            pdf.set_text_color(*_TEXT_MUTED)
            pdf.set_x(10)
            pdf.multi_cell(
                190, 4,
                "* Prices and NAVs are indicative. Actual execution price will be the prevailing "
                "market price or NAV at time of transaction. Insurance premiums are annual. "
                "Items marked * are IRDAI-regulated products outside SEBI purview.",
                align="L",
            )
        else:
            pdf.set_font("helvetica", "I", 9)
            pdf.set_text_color(*_TEXT_MUTED)
            pdf.cell(0, 10, "No recommendations attached to this advice note.", ln=True, align="C")
        pdf.ln(6)

        # ══════════════════ SECTION E - Rationale ══════════════════
        if recommendations and any(r.get("rationale") for r in recommendations):
            if pdf.get_y() > 230:
                pdf.add_page()
            InvestmentAdviceNotePDF._section_header(pdf, "Section E - Rationale for Advice")

            for rec in recommendations:
                rationale = rec.get("rationale", "")
                if rationale:
                    name = rec.get("product_name", "Product")
                    action = rec.get("action", "")
                    pdf.set_x(10)
                    pdf.set_font("helvetica", "B", 8)
                    pdf.set_text_color(*_NAVY)
                    
                    name_txt = f" {name}"
                    font_size = 8.0
                    width = pdf.get_string_width(name_txt)
                    max_allowed_w = 40  # 42mm total width, 2mm padding safety
                    
                    while width > max_allowed_w and font_size > 6.0:
                        font_size -= 0.5
                        pdf.set_font("helvetica", "B", font_size)
                        width = pdf.get_string_width(name_txt)
                        
                    if width > max_allowed_w:
                        while len(name_txt) > 4 and width > max_allowed_w:
                            name_txt = name_txt[:-4] + "..."
                            width = pdf.get_string_width(name_txt)
                            
                    pdf.cell(42, 6, name_txt, border="LT")
                    pdf.set_font("helvetica", "", 7.5)
                    pdf.set_text_color(*_TEXT_DARK)
                    pdf.multi_cell(148, 6, f" {action}. {rationale}", border="RT")
                    # Bottom border
                    pdf.set_x(10)
                    pdf.cell(190, 0, "", border="T")
                    pdf.ln(1)
            pdf.ln(6)

        # ══════════════════ SECTION F - Risk Disclosures ══════════════════
        if pdf.get_y() > 230:
            pdf.add_page()
        InvestmentAdviceNotePDF._section_header(pdf, "Section F - Risk Disclosures")

        risk_disclosures = [
            ("Market / Price Risk", "Equity and ETF investments are subject to market fluctuations. Past performance is not indicative of future returns. The value of investments may fall below the invested amount."),
            ("Mutual Fund Risk", "Mutual Fund investments are subject to market risks. Please read all scheme documents (SID and KIM) carefully before investing. NAV may go up or down depending on market conditions."),
            ("Interest Rate / Duration Risk", "SBI Magnum Medium Duration Fund carries interest rate duration risk. A rise in interest rates will negatively affect NAV. Suitable only for investors with a minimum 2-3 year horizon."),
            ("Gold / Commodity Risk", "Gold Bees ETF tracks domestic gold prices influenced by global commodity prices, INR/USD exchange rates and geopolitical factors. Gold does not generate any income (no dividend or interest)."),
            ("Concentration / Stock Risk", "Individual equity positions carry stock-specific risk including regulatory action, management changes, sector headwinds and liquidity events. Portfolio diversification is recommended."),
        ]

        has_insurance = recommendations and any(
            rec.get("product_type", "").lower() in ("life-insurance", "life_insurance") 
            for rec in recommendations
        )
        if has_insurance:
            risk_disclosures.append((
                "Life Insurance — IRDAI Regulated *",
                "Life insurance is regulated by IRDAI, NOT SEBI. SEBI has no jurisdiction over this product. Any grievance relating to insurance advice must be directed to IRDAI (www.irdai.gov.in). The IA's advisory services for this product are outside SEBI's regulatory purview and no recourse is available from SEBI or IAASB. Client has signed a separate non-SEBI disclosure and declaration."
            ))

        for label, text in risk_disclosures:
            if pdf.get_y() > 265:
                pdf.add_page()
            pdf.set_x(10)
            
            is_insurance = label.startswith("Life Insurance")
            
            if is_insurance:
                pdf.set_fill_color(254, 236, 236)  # Pink background for insurance key cell
                pdf.set_font("helvetica", "B", 7.5)
                pdf.set_text_color(180, 60, 60)   # Red text color
                pdf.cell(38, 6, f" {label}", border="LTB", fill=True)
            else:
                pdf.set_font("helvetica", "B", 7.5)
                pdf.set_text_color(*_NAVY)
                pdf.cell(38, 6, f" {label}", border="LTB")
                
            pdf.set_font("helvetica", "", 7)
            pdf.set_text_color(*_TEXT_DARK)
            pdf.multi_cell(152, 6, f" {text}", border="RTB")
        pdf.ln(6)

        # ══════════════════ SECTION G - Disclosures ══════════════════
        if pdf.get_y() > 210:
            pdf.add_page()
        InvestmentAdviceNotePDF._section_header(
            pdf, "Section G - Conflict of Interest and AI Usage Disclosure"
        )

        disclosure_fields = [
            ("Conflict of Interest", note_data.get("conflict_of_interest_text", "No conflicts of interest declared.")),
            ("No Execution by IA", note_data.get("no_execution_text", "The IA is not authorised to execute trades on behalf of the client.")),
            ("AI Tool Disclosure", note_data.get("ai_usage_text", "No AI tools were used in the preparation of this advice note.")),
        ]

        for label, text in disclosure_fields:
            if pdf.get_y() > 260:
                pdf.add_page()
            pdf.set_x(10)
            pdf.set_font("helvetica", "B", 7.5)
            pdf.set_text_color(*_NAVY)
            pdf.cell(38, 6, f" {label}", border="LTB")
            pdf.set_font("helvetica", "", 7)
            pdf.set_text_color(*_TEXT_DARK)
            pdf.multi_cell(152, 6, f" {text}", border="RTB")
        pdf.ln(8)

        # ══════════════════ DISCLAIMER ══════════════════
        if pdf.get_y() > 220:
            pdf.add_page()
        pdf.set_font("helvetica", "B", 8)
        pdf.set_text_color(*_RED_MUTED)
        pdf.cell(0, 6, "IMPORTANT DISCLAIMER", ln=True, align="C")
        pdf.set_font("helvetica", "I", 6.5)
        pdf.set_text_color(*_TEXT_MUTED)
        pdf.multi_cell(0, 4, _SEBI_DISCLAIMER, align="C")
        pdf.ln(3)

        # ══════════════════ SIGNATURES ══════════════════
        if pdf.get_y() > 240:
            pdf.add_page()

        sig_y = pdf.get_y() + 5
        pdf.set_y(sig_y)

        # Left: IA Signature
        pdf.set_x(10)
        pdf.set_font("helvetica", "B", 8)
        pdf.set_text_color(*_NAVY)
        entity = entity_name or advisor_name or "Investment Advisor"
        pdf.cell(90, 5, f"For {entity}", ln=True)
        pdf.set_x(10)
        pdf.ln(12)
        pdf.set_x(10)
        pdf.set_font("helvetica", "", 7)
        pdf.set_text_color(*_TEXT_DARK)
        pdf.cell(90, 4, "________________________________", ln=True)
        pdf.set_x(10)
        po_name = note_data.get("principal_officer_name", "Principal Officer")
        pdf.cell(90, 5, po_name, ln=True)
        pdf.set_x(10)
        pdf.set_font("helvetica", "I", 7)
        pdf.set_text_color(*_TEXT_MUTED)
        pdf.cell(90, 4, f"Principal Officer  |  Reg No: {ia_reg_no}", ln=True)
        pdf.set_x(10)
        pdf.cell(90, 4, f"Date: {note_data.get('date_of_issue', 'N/A')}", ln=True)

        # Right: Client Acknowledgement
        client_sig_y = sig_y
        pdf.set_xy(110, client_sig_y)
        pdf.set_font("helvetica", "B", 8)
        pdf.set_text_color(*_NAVY)
        pdf.cell(90, 5, "Client Acknowledgement", ln=True)
        pdf.set_xy(110, client_sig_y + 17)
        pdf.set_font("helvetica", "", 7)
        pdf.set_text_color(*_TEXT_DARK)
        pdf.cell(90, 4, "________________________________", ln=True)
        pdf.set_xy(110, pdf.get_y())
        pdf.cell(90, 5, "I have read and understood this", ln=True)
        pdf.set_xy(110, pdf.get_y())
        pdf.cell(90, 5, "Investment Advice Note.", ln=True)
        pdf.set_xy(110, pdf.get_y())
        pdf.set_font("helvetica", "B", 7)
        pdf.cell(90, 5, client.get("client_name", "Client"), ln=True)
        pdf.set_xy(110, pdf.get_y())
        pdf.set_font("helvetica", "I", 7)
        pdf.set_text_color(*_TEXT_MUTED)
        pdf.cell(90, 4, f"Client ID: {client.get('client_code', 'N/A')}", ln=True)
        pdf.set_xy(110, pdf.get_y())
        pdf.cell(90, 4, "Date: ___________________", ln=True)

        # Record retention footer
        pdf.ln(8)
        pdf.set_font("helvetica", "I", 6)
        pdf.set_text_color(*_TEXT_MUTED)
        pdf.multi_cell(0, 4, _RECORD_RETENTION, align="C")

        return bytes(pdf.output())


# ════════════════════════════════════════════════════════════════════
# DOCX Generator
# ════════════════════════════════════════════════════════════════════

class InvestmentAdviceNoteDOCX:
    """Generates a SEBI-compliant Investment Advice Note Word document."""

    @staticmethod
    def _style_cell(
        cell, 
        bg_color=None, 
        border_color="CCCCCC", 
        top=True, 
        bottom=True, 
        left=True, 
        right=True, 
        top_pad=80, 
        bottom_pad=80, 
        left_pad=120, 
        right_pad=120
    ):
        """Apply shading, custom borders, cell margins, and paragraph spacing to a cell."""
        tcPr = cell._tc.get_or_add_tcPr()
        
        # 1. Shading (Background Color)
        if bg_color:
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}"/>')
            tcPr.append(shd)
            
        # 2. Borders
        borders_str = f'<w:tcBorders {nsdecls("w")}>'
        borders_str += f'<w:top w:val="{"single" if top else "nil"}" w:sz="4" w:space="0" w:color="{border_color}"/>'
        borders_str += f'<w:bottom w:val="{"single" if bottom else "nil"}" w:sz="4" w:space="0" w:color="{border_color}"/>'
        borders_str += f'<w:left w:val="{"single" if left else "nil"}" w:sz="4" w:space="0" w:color="{border_color}"/>'
        borders_str += f'<w:right w:val="{"single" if right else "nil"}" w:sz="4" w:space="0" w:color="{border_color}"/>'
        borders_str += '</w:tcBorders>'
        tcPr.append(parse_xml(borders_str))
        
        # 3. Cell margins (padding)
        tcMar = parse_xml(
            f'<w:tcMar {nsdecls("w")}>'
            f'<w:top w:w="{top_pad}" w:type="dxa"/>'
            f'<w:bottom w:w="{bottom_pad}" w:type="dxa"/>'
            f'<w:left w:w="{left_pad}" w:type="dxa"/>'
            f'<w:right w:w="{right_pad}" w:type="dxa"/>'
            f'</w:tcMar>'
        )
        tcPr.append(tcMar)

        # 4. Paragraph formatting for all paragraphs inside cell
        for p in cell.paragraphs:
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)

    @staticmethod
    def _add_section_heading(doc, text: str):
        """Add a professional styled section heading with dark blue color."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(0x1F, 0x51, 0x88)
        return p

    @staticmethod
    def _add_kv_grid_table(doc, fields: List[tuple]):
        """Add a 4-column key-value grid table."""
        rows_data = []
        i = 0
        while i < len(fields):
            f1 = fields[i]
            is_full_width_1 = len(f1) > 2 and f1[2]
            
            if is_full_width_1:
                rows_data.append((f1[0], f1[1], None, None))
                i += 1
            else:
                if i + 1 < len(fields):
                    f2 = fields[i+1]
                    is_full_width_2 = len(f2) > 2 and f2[2]
                    if is_full_width_2:
                        rows_data.append((f1[0], f1[1], "", ""))
                        i += 1
                    else:
                        rows_data.append((f1[0], f1[1], f2[0], f2[1]))
                        i += 2
                else:
                    rows_data.append((f1[0], f1[1], "", ""))
                    i += 1

        # Add table
        table = doc.add_table(rows=len(rows_data), cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Column widths: 
        # Key 1: 1.2 inches, Val 1: 2.3 inches, Key 2: 1.2 inches, Val 2: 2.3 inches
        col_widths = [Inches(1.2), Inches(2.3), Inches(1.2), Inches(2.3)]
        
        for r_idx, (k1, v1, k2, v2) in enumerate(rows_data):
            row = table.rows[r_idx]
            
            if k2 is None and v2 is None:
                # Full width row (merge cells 1, 2, and 3)
                cell_k = row.cells[0]
                cell_v = row.cells[1]
                cell_v.merge(row.cells[2]).merge(row.cells[3])
                
                # Key style
                InvestmentAdviceNoteDOCX._style_cell(cell_k, bg_color="F4F6F9", border_color="D3D3D3")
                pk = cell_k.paragraphs[0]
                run_k = pk.add_run(str(k1))
                run_k.bold = True
                run_k.font.size = Pt(9)
                run_k.font.name = 'Arial'
                run_k.font.color.rgb = RGBColor(0x1F, 0x51, 0x88)
                
                # Val style
                InvestmentAdviceNoteDOCX._style_cell(cell_v, bg_color="FFFFFF", border_color="D3D3D3")
                pv = cell_v.paragraphs[0]
                run_v = pv.add_run(str(v1 or "N/A"))
                run_v.font.size = Pt(9)
                run_v.font.name = 'Arial'
            else:
                # Key 1
                cell_k1 = row.cells[0]
                InvestmentAdviceNoteDOCX._style_cell(cell_k1, bg_color="F4F6F9", border_color="D3D3D3")
                pk1 = cell_k1.paragraphs[0]
                run_k1 = pk1.add_run(str(k1))
                run_k1.bold = True
                run_k1.font.size = Pt(9)
                run_k1.font.name = 'Arial'
                run_k1.font.color.rgb = RGBColor(0x1F, 0x51, 0x88)
                
                # Val 1
                cell_v1 = row.cells[1]
                InvestmentAdviceNoteDOCX._style_cell(cell_v1, bg_color="FFFFFF", border_color="D3D3D3")
                pv1 = cell_v1.paragraphs[0]
                run_v1 = pv1.add_run(str(v1 or "N/A"))
                run_v1.font.size = Pt(9)
                run_v1.font.name = 'Arial'
                
                # Key 2
                cell_k2 = row.cells[2]
                InvestmentAdviceNoteDOCX._style_cell(cell_k2, bg_color="F4F6F9", border_color="D3D3D3")
                pk2 = cell_k2.paragraphs[0]
                run_k2 = pk2.add_run(str(k2))
                run_k2.bold = True
                run_k2.font.size = Pt(9)
                run_k2.font.name = 'Arial'
                run_k2.font.color.rgb = RGBColor(0x1F, 0x51, 0x88)
                
                # Val 2
                cell_v2 = row.cells[3]
                InvestmentAdviceNoteDOCX._style_cell(cell_v2, bg_color="FFFFFF", border_color="D3D3D3")
                pv2 = cell_v2.paragraphs[0]
                run_v2 = pv2.add_run(str(v2 or "N/A"))
                run_v2.font.size = Pt(9)
                run_v2.font.name = 'Arial'

        # Set widths on all unmerged cells
        for row in table.rows:
            if len(row.cells) == 4:
                for idx, w in enumerate(col_widths):
                    row.cells[idx].width = w

    @staticmethod
    def _add_two_column_table(doc, fields: List[tuple]):
        """Add a 2-column table (for Section E, F, G)."""
        table = doc.add_table(rows=len(fields), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Column widths: Left: 1.8 inches, Right: 5.2 inches (Total = 7.0 inches)
        col_widths = [Inches(1.8), Inches(5.2)]
        
        for r_idx, (k, v) in enumerate(fields):
            row = table.rows[r_idx]
            
            # Left key cell
            cell_k = row.cells[0]
            is_insurance = str(k).startswith("Life Insurance")
            
            if is_insurance:
                # Pink background for insurance key cell
                InvestmentAdviceNoteDOCX._style_cell(cell_k, bg_color="FEECEC", border_color="D3D3D3")
                pk = cell_k.paragraphs[0]
                run_k = pk.add_run(str(k))
                run_k.bold = True
                run_k.font.size = Pt(9)
                run_k.font.name = 'Arial'
                run_k.font.color.rgb = RGBColor(180, 60, 60)  # Red text
            else:
                InvestmentAdviceNoteDOCX._style_cell(cell_k, bg_color="F4F6F9", border_color="D3D3D3")
                pk = cell_k.paragraphs[0]
                run_k = pk.add_run(str(k))
                run_k.bold = True
                run_k.font.size = Pt(9)
                run_k.font.name = 'Arial'
                run_k.font.color.rgb = RGBColor(0x1F, 0x51, 0x88)
            
            # Right value cell
            cell_v = row.cells[1]
            InvestmentAdviceNoteDOCX._style_cell(cell_v, bg_color="FFFFFF", border_color="D3D3D3")
            pv = cell_v.paragraphs[0]
            run_v = pv.add_run(str(v or "N/A"))
            run_v.font.size = Pt(9)
            run_v.font.name = 'Arial'
            
        # Set widths
        for row in table.rows:
            if len(row.cells) == 2:
                row.cells[0].width = col_widths[0]
                row.cells[1].width = col_widths[1]

    @staticmethod
    def _add_recommendation_table(doc, recommendations: List[dict], date_of_issue: str = ""):
        """Add the recommendations table with styled columns, color-coded actions, and merged footnote (Section D)."""
        headers = [
            "#", 
            "Product / Scheme Name", 
            "ISIN / Scheme Code", 
            "Product Type", 
            "Action", 
            "Amount / Units", 
            "Indicative Price / NAV (Rs.)/Premium"
        ]
        
        # 1 header row + N data rows + 1 footnote row
        total_rows = 1 + len(recommendations) + 1
        table = doc.add_table(rows=total_rows, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 1. Style Header Row (Row 0)
        hdr_row = table.rows[0]
        for i, h in enumerate(headers):
            cell = hdr_row.cells[i]
            InvestmentAdviceNoteDOCX._style_cell(cell, bg_color="1F5188", border_color="1F5188", top_pad=100, bottom_pad=100)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            
            # Header Alignments
            if h == "#" or h == "Action" or "ISIN" in h or "Type" in h:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif "Price" in h:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
            run = p.add_run(h)
            run.bold = True
            run.font.size = Pt(8.5)
            run.font.name = 'Arial'
            run.font.color.rgb = RGBColor(255, 255, 255)
            
        # Data Rows (Row 1 to N)
        for idx, rec in enumerate(recommendations):
            row = table.rows[idx + 1]
            bg_color = "FFFFFF" if idx % 2 == 0 else "F4F6F9"
            
            p_name = rec.get("product_name", "N/A")
            isin = rec.get("isin_code_scheme_code_uin", "") or "N/A"
            p_type = rec.get("product_type", "").replace("_", " ").title()
            action = rec.get("action", "BUY")
            amt_units = format_amount_units_python(rec)
            
            price_val = rec.get("indicative_price_nav")
            if price_val:
                try:
                    price_str = f"{float(price_val):,.2f}"
                except Exception:
                    price_str = str(price_val)
            else:
                price_str = "N/A"
                
            # Values array
            values = [str(idx + 1), p_name, isin, p_type, action, amt_units, price_str]
            
            for i, val in enumerate(values):
                cell = row.cells[i]
                InvestmentAdviceNoteDOCX._style_cell(cell, bg_color=bg_color, border_color="D3D3D3")
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)
                
                # Alignments
                if i in (0, 2, 3, 4):
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif i == 6:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                run = p.add_run(val)
                run.font.size = Pt(8.5)
                run.font.name = 'Arial'
                
                # Special colors for Action column
                if i == 4:
                    run.bold = True
                    if val == "BUY":
                        run.font.color.rgb = RGBColor(0, 128, 0)
                    elif val == "HOLD":
                        run.font.color.rgb = RGBColor(0, 70, 160)
                    elif val == "SELL":
                        run.font.color.rgb = RGBColor(180, 60, 60)
                    elif val == "REVIEW":
                        run.font.color.rgb = RGBColor(139, 69, 19)
                        
        # Footnote Row (Row N+1)
        ft_row = table.rows[len(recommendations) + 1]
        ft_cell = ft_row.cells[0]
        for c in range(1, len(headers)):
            ft_cell.merge(ft_row.cells[c])
            
        InvestmentAdviceNoteDOCX._style_cell(ft_cell, bg_color="EBF8EB", border_color="D3D3D3", top_pad=80, bottom_pad=80)
        p = ft_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        
        issue_date_str = date_of_issue or "______________"
        footnote_text = (
            f"* Prices and NAVs are indicative as of {issue_date_str}. Actual execution price will be "
            "the prevailing market price or NAV at time of transaction. Insurance premium is annual. "
            "Items marked * are IRDAI-regulated products outside SEBI purview — see Section F."
        )
        run = p.add_run(footnote_text)
        run.italic = True
        run.font.size = Pt(8)
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(46, 125, 50)
        
        # Set column widths (Total = 7.0 inches)
        col_widths = [Inches(0.35), Inches(2.0), Inches(1.15), Inches(0.9), Inches(0.6), Inches(1.1), Inches(0.9)]
        for row in table.rows:
            if len(row.cells) == len(headers):
                for idx, w in enumerate(col_widths):
                    row.cells[idx].width = w

    @staticmethod
    def generate_docx(
        note_data: dict,
        ia_data: Optional[dict] = None,
    ) -> io.BytesIO:
        """
        Generate a SEBI Investment Advice Note as a beautifully styled Word document.
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is not installed on this system.")

        doc = Document()
        
        # Set Page Margins to 0.75" on all sides
        section = doc.sections[0]
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

        # Set Normal text style defaults
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(9.5)

        client = note_data.get("client_snapshot", {})
        recommendations = note_data.get("recommendations", [])

        advisor_name = ia_data.get("name_of_ia", "") if ia_data else ""
        entity_name = ia_data.get("name_of_entity", "") if ia_data else ""
        ia_reg_no = ia_data.get("ia_registration_number", "") if ia_data else ""

        # ── Header ──
        header = section.header
        h_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        h_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        h_para.text = ""
        h_run = h_para.add_run(f"{entity_name or advisor_name}  |  SEBI Reg: {ia_reg_no}")
        h_run.font.size = Pt(8)
        h_run.font.name = 'Arial'
        h_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

        # ── Document Title ──
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.paragraph_format.space_before = Pt(12)
        title_para.paragraph_format.space_after = Pt(2)
        t_run = title_para.add_run("INVESTMENT ADVICE NOTE")
        t_run.bold = True
        t_run.font.size = Pt(18)
        t_run.font.name = 'Arial'
        t_run.font.color.rgb = RGBColor(0x1F, 0x51, 0x88)

        sub_para = doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_para.paragraph_format.space_after = Pt(18)
        s_run = sub_para.add_run("Issued under SEBI (Investment Advisers) Regulations, 2013  |  Regulation 16 & 17")
        s_run.font.size = Pt(9)
        s_run.font.name = 'Arial'
        s_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        # ══════════════════ SECTION A ══════════════════
        basl_id = ia_data.get("basl_membership_id", "") if ia_data else ""
        website_val = ia_data.get("website", "") if ia_data else ""
        addr = ia_data.get("registered_address", "N/A") if ia_data else "N/A"
        po_name = note_data.get("principal_officer_name", "N/A")
        po_reg = note_data.get("principal_officer_reg_no", "")
        po_display = f"{po_name}, Reg No: {po_reg}" if po_reg else po_name
        issue_date_formatted = format_date_issue(note_data.get("date_of_issue"))

        fields_a = [
            ("IA / Firm Name", entity_name or advisor_name),
            ("Advice Note No.", note_data.get("advice_note_no", "N/A")),
            ("SEBI Registration No.", ia_reg_no),
            ("Date of Issue", issue_date_formatted),
            ("IAASB", f"BSE Limited (IAASB) - {basl_id}" if basl_id else "BSE Limited (IAASB)"),
            ("Advice Validity", note_data.get("advice_validity_custom_text", "N/A")),
            ("Registered Address", addr),
            ("Principal Officer", po_display),
            ("Website", website_val or "N/A"),
            ("Advice Category", note_data.get("advice_category", "Comprehensive Advisory")),
        ]
        InvestmentAdviceNoteDOCX._add_section_heading(doc, "Section A — Investment Adviser Details")
        InvestmentAdviceNoteDOCX._add_kv_grid_table(doc, fields_a)
        doc.add_paragraph()

        # ══════════════════ SECTION B ══════════════════
        risk_score = client.get("risk_profile_score")
        risk_profile_str = client.get("risk_profile", "N/A")
        if risk_score and risk_score != "N/A":
            risk_profile_str = f"{risk_profile_str} (Risk Score: {risk_score} / 100)"

        liabilities_val = client.get("existing_liabilities")
        if liabilities_val is not None:
            try:
                liabilities_str = f"Rs. {float(liabilities_val):,.0f}"
            except Exception:
                liabilities_str = str(liabilities_val)
        else:
            liabilities_str = "N/A"

        # Email & Mobile
        email = client.get("email") or ""
        phone = client.get("phone_number") or ""
        email_mobile = f"{email} | {phone}" if email and phone else (email or phone or "N/A")

        # DOB & Age
        dob_val = client.get("date_of_birth", "")
        dob_formatted = format_dob(dob_val)
        age = calculate_age(dob_val)
        dob_display = f"{dob_formatted} (Age: {age} years)" if age else dob_formatted

        # Fee Mode and Amount
        fee_mode = note_data.get("fee_mode", "N/A").replace("_", " ").title()
        fee_amount = float(note_data.get("fee_amount", 0))
        fee_display = f"{fee_mode} | Rs. {fee_amount:,.0f}" if fee_amount > 0 else fee_mode

        # Assets Under Advice
        aua = float(note_data.get("assets_under_advice", 0))
        aua_display = f"Rs. {aua:,.0f} (approx.)" if aua > 0 else "N/A"

        # Recommended Asset Allocation
        rec_alloc = note_data.get("recommended_asset_allocation")
        alloc_text = "N/A"
        sub_allocs = {}
        if rec_alloc:
            if isinstance(rec_alloc, str):
                try:
                    import json as _json
                    rec_alloc = _json.loads(rec_alloc)
                except Exception:
                    pass
            if isinstance(rec_alloc, dict):
                sub_allocs = rec_alloc.get("sub_assets") or {}
                alloc_parts = []
                for k in ("Debt", "Equity", "Commodities"):
                    if k in rec_alloc:
                        alloc_parts.append(f"{k}: {rec_alloc[k]}%")
                for k, v in rec_alloc.items():
                    if k not in ("Debt", "Equity", "Commodities", "sub_assets"):
                        alloc_parts.append(f"{k}: {v}%")
                alloc_text = "  |  ".join(alloc_parts) if alloc_parts else "N/A"
            elif isinstance(rec_alloc, str) and rec_alloc != "{}":
                alloc_text = rec_alloc

        if sub_allocs:
            label_map = {
                "fixed_deposits_bonds_percentage": "Fixed Deposits / Bonds",
                "mutual_fund_debt_percentage": "Debt Mutual Funds",
                "ulip_debt_percentage": "Debt ULIPs",
                "etf_debt_percentage": "Debt ETFs",
                "stocks_percentage": "Direct Equity (Stocks)",
                "mutual_fund_equity_percentage": "Equity Mutual Funds",
                "ulip_equity_percentage": "Equity ULIPs",
                "etf_equity_percentage": "Equity ETFs",
                "gold_etf_percentage": "Gold ETFs",
                "silver_etf_percentage": "Silver ETFs",
                "etf_commodity_percentage": "Commodity ETFs"
            }
            sub_parts = []
            for k, val in sub_allocs.items():
                try:
                    num_val = float(val) if val is not None else 0.0
                except Exception:
                    num_val = 0.0
                if num_val > 0:
                    friendly_name = label_map.get(k, k.replace("_percentage", "").replace("_", " ").title())
                    sub_parts.append(f"{friendly_name}: {num_val}%")
            if sub_parts:
                alloc_text += "\nSub-Asset Breakdown:\n" + "\n".join(f"• {x}" for x in sub_parts)

        fields_b = [
            ("Client Full Name", client.get("client_name", "N/A")),
            ("Client ID", client.get("client_code", "N/A")),
            ("PAN Number", client.get("pan_number", "N/A")),
            ("Date of Birth", dob_display),
            ("Address", client.get("address", "N/A")),
            ("Email / Mobile", email_mobile),
            ("Risk Profile Category", risk_profile_str),
            ("Risk Profiling Date", format_dob(client.get("risk_profile_date"))),
            ("Investment Horizon", client.get("investment_horizon", "N/A")),
            ("Annual Income Band", note_data.get("annual_income_band", "N/A")),
            ("Assets Under Advice", aua_display),
            ("Fee Mode and Amount", fee_display),
            ("Recommended Asset Allocation", alloc_text, True),
            ("Date of Allocation", format_dob(note_data.get("date_of_allocation"))),
        ]
        InvestmentAdviceNoteDOCX._add_section_heading(doc, "Section B — Client Details and Risk Profile")
        InvestmentAdviceNoteDOCX._add_kv_grid_table(doc, fields_b)
        doc.add_paragraph()

        # ══════════════════ SECTION C ══════════════════
        fields_c = [
            ("Advice Suitable?", note_data.get("suitability_assessment", "N/A")),
            ("Date of Allocation", format_dob(note_data.get("date_of_allocation"))),
            ("Suitability Basis", note_data.get("suitability_basis", "N/A"), True),
            ("Investor Advice", note_data.get("investor_advice", "N/A"), True),
            ("Current Asset Allocation", note_data.get("current_asset_allocation", "N/A"), True),
            ("Rebalancing Rationale", note_data.get("rebalancing_rationale", "N/A"), True),
        ]
        InvestmentAdviceNoteDOCX._add_section_heading(doc, "Section C — Suitability Assessment")
        InvestmentAdviceNoteDOCX._add_kv_grid_table(doc, fields_c)
        doc.add_paragraph()

        # ══════════════════ SECTION D ══════════════════
        InvestmentAdviceNoteDOCX._add_section_heading(doc, "Section D — Investment Recommendations [Regulation 16]")
        if recommendations:
            InvestmentAdviceNoteDOCX._add_recommendation_table(doc, recommendations, date_of_issue=issue_date_formatted)
        else:
            doc.add_paragraph("No recommendations attached to this advice note.").italic = True
        doc.add_paragraph()

        # ══════════════════ SECTION E ══════════════════
        if recommendations and any(r.get("rationale") for r in recommendations):
            rationale_rows = []
            r_idx = 1
            for rec in recommendations:
                if rec.get("rationale"):
                    prod_name = rec.get("product_name", "Product")
                    rationale_text = f"{rec.get('action', '')}. {rec.get('rationale', '')}"
                    rationale_rows.append((f"{r_idx}. {prod_name}", rationale_text))
                    r_idx += 1
            
            if rationale_rows:
                InvestmentAdviceNoteDOCX._add_section_heading(doc, "Section E — Rationale for Advice")
                InvestmentAdviceNoteDOCX._add_two_column_table(
                    doc, 
                    rationale_rows
                )
                doc.add_paragraph()

        # ══════════════════ SECTION F ══════════════════
        fields_f = [
            ("Market / Price Risk", "Equity and ETF investments are subject to market fluctuations. Past performance is not indicative of future returns. The value of investments may fall below the invested amount."),
            ("Mutual Fund Risk", "Mutual Fund investments are subject to market risks. Please read all scheme documents (SID and KIM) carefully before investing. NAV may go up or down depending on market conditions."),
            ("Interest Rate / Duration Risk", "SBI Magnum Medium Duration Fund carries interest rate duration risk. A rise in interest rates will negatively affect NAV. Suitable only for investors with a minimum 2-3 year horizon."),
            ("Gold / Commodity Risk", "Gold Bees ETF tracks domestic gold prices influenced by global commodity prices, INR/USD exchange rates and geopolitical factors. Gold does not generate any income (no dividend or interest)."),
            ("Concentration / Stock Risk", "Individual equity positions carry stock-specific risk including regulatory action, management changes, sector headwinds and liquidity events. Portfolio diversification is recommended."),
        ]
        
        has_insurance = recommendations and any(
            rec.get("product_type", "").lower() in ("life-insurance", "life_insurance") 
            for rec in recommendations
        )
        if has_insurance:
            fields_f.append((
                "Life Insurance — IRDAI Regulated *",
                "Life insurance is regulated by IRDAI, NOT SEBI. SEBI has no jurisdiction over this product. Any grievance relating to insurance advice must be directed to IRDAI (www.irdai.gov.in). The IA's advisory services for this product are outside SEBI's regulatory purview and no recourse is available from SEBI or IAASB. Client has signed a separate non-SEBI disclosure and declaration."
            ))

        InvestmentAdviceNoteDOCX._add_section_heading(doc, "Section F — Risk Disclosures")
        InvestmentAdviceNoteDOCX._add_two_column_table(
            doc, 
            fields_f
        )
        doc.add_paragraph()

        # ══════════════════ SECTION G ══════════════════
        fields_g = [
            ("Conflict of Interest", note_data.get("conflict_of_interest_text", "No conflicts of interest declared.")),
            ("No Execution by IA", note_data.get("no_execution_text", "The IA is not authorised to execute trades on behalf of the client.")),
            ("AI Tool Disclosure", note_data.get("ai_usage_text", "No AI tools were used in the preparation of this advice note.")),
        ]
        InvestmentAdviceNoteDOCX._add_section_heading(
            doc, 
            "Section G — Conflict of Interest and AI Usage Disclosure [Reg. 18 and 15(14)]"
        )
        InvestmentAdviceNoteDOCX._add_two_column_table(
            doc, 
            fields_g
        )
        doc.add_paragraph()

        # ══════════════════ DISCLAIMER ══════════════════
        doc.add_heading("IMPORTANT DISCLAIMER", level=2)
        disclaimer_para = doc.add_paragraph(_SEBI_DISCLAIMER)
        disclaimer_para.style = doc.styles["Intense Quote"]
        for run in disclaimer_para.runs:
            run.font.size = Pt(8.5)
            run.font.name = 'Arial'
            run.font.color.rgb = RGBColor(0x1F, 0x51, 0x88)

        # ══════════════════ SIGNATURES ══════════════════
        doc.add_paragraph()
        sig_table = doc.add_table(rows=4, cols=2)
        
        # Left: IA
        sig_table.cell(0, 0).text = f"For {entity_name or advisor_name or 'Investment Advisor'}"
        sig_table.cell(1, 0).text = "\n\n__________________________"
        po_name = note_data.get("principal_officer_name", "Principal Officer")
        sig_table.cell(2, 0).text = f"{po_name}\nPrincipal Officer\nSEBI Reg. No.: {ia_reg_no}"
        sig_table.cell(3, 0).text = f"Date: {issue_date_formatted}"
        
        # Right: Client
        sig_table.cell(0, 1).text = "Client Acknowledgement"
        sig_table.cell(1, 1).text = "\n\n__________________________"
        sig_table.cell(2, 1).text = (
            "I have read and understood this Investment Advice Note "
            "and the risks disclosed herein."
        )
        sig_table.cell(3, 1).text = (
            f"{client.get('client_name', 'Client')}\n"
            f"Client ID: {client.get('client_code', 'N/A')}\n"
            f"Date: ___________________"
        )

        # Set column widths & style signature cells to be borderless and padded
        for row in sig_table.rows:
            row.cells[0].width = Inches(3.5)
            row.cells[1].width = Inches(3.5)
            for cell in row.cells:
                InvestmentAdviceNoteDOCX._style_cell(
                    cell, 
                    bg_color="FFFFFF", 
                    top=False, 
                    bottom=False, 
                    left=False, 
                    right=False,
                    top_pad=60,
                    bottom_pad=60,
                    left_pad=100,
                    right_pad=100
                )
                
                # Make text runs 9 pt Arial
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)
                        run.font.name = 'Arial'

        # Record retention
        doc.add_paragraph()
        retention = doc.add_paragraph(_RECORD_RETENTION)
        retention.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in retention.runs:
            run.font.size = Pt(7)
            run.italic = True
            run.font.name = 'Arial'
            run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        # ── Footer ──
        footer = section.footer
        f_p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        f_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        f_p.text = ""
        f_run = f_p.add_run(
            f"Prepared by: {advisor_name}  |  Entity: {entity_name}  |  Reg No: {ia_reg_no}"
        )
        f_run.font.size = Pt(7)
        f_run.font.name = 'Arial'
        f_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        f_run.italic = True

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
